"""docx exports (python-docx) for household-services results.

Formats the engine's typed result into a Word report. Does no math. Returns the
path written.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from engine.lhhs import LHHSResult


def _money(value: float) -> str:
    return f"${value:,.0f}"


def _sources_appendix(doc, sources: list | None) -> None:
    """Append the raw look-up data behind the numbers as a report appendix."""
    if not sources:
        return
    doc.add_page_break()
    doc.add_heading("Appendix: raw data and sources", level=1)
    for src in sources:
        doc.add_heading(src.get("title", ""), level=2)
        if src.get("citation"):
            doc.add_paragraph(src["citation"]).italic = True
        cols = src.get("columns") or []
        rows = src.get("rows") or []
        if cols and rows:
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            for i, label in enumerate(cols):
                table.rows[0].cells[i].text = str(label)
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)


def lhhs_report(result: LHHSResult, path: str, *, heading: str | None = None,
                sources: list | None = None) -> str:
    doc = Document()
    doc.add_heading(heading or "Loss of Household Services", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Values use the replacement-cost method (Determining Economic Damages, "
        "Chapter 6), with annual household-production values from the Expectancy "
        "Data Dollar Value of a Day tables, grown for wage inflation and "
        "discounted to present value."
    ).font.size = Pt(10)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(
        ["Year", "Annual value (grown)", "Loss %", "Annual loss", "Present value"]
    ):
        hdr[i].text = label

    for r in result.rows:
        cells = table.add_row().cells
        cells[0].text = str(r.year)
        cells[1].text = _money(r.base_annual_value)
        cells[2].text = f"{r.loss_percent * 100:.0f}%"
        cells[3].text = _money(r.annual_loss)
        cells[4].text = _money(r.present_value)

    doc.add_paragraph()
    for label, value in (
        ("Past present value", result.past_present_value),
        ("Future present value", result.future_present_value),
        ("Total present value", result.total_present_value),
    ):
        para = doc.add_paragraph()
        run = para.add_run(f"{label}: {value:,.2f}")
        if label.startswith("Total"):
            run.bold = True

    _sources_appendix(doc, sources)
    doc.save(path)
    return path
